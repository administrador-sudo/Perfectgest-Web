# -*- coding: utf-8 -*-
"""Gera modelo Word do PPP (Anexo XVII, IN PRES/INSS 128/2022). Campos em branco."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "PPP_INSS_Modelo_Atual.docx"
NAVY = RGBColor(0x1A, 0x36, 0x5D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x4B, 0x55, 0x63)
LINE = "________________________________________________"


def set_run(run, size=10, bold=False, color=BLACK, name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def p_style(p, before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT, space=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = space
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = align


def add_text(p, text, **kw):
    r = p.add_run(text)
    set_run(r, **kw)
    return r


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, color="1A365D", sz="8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, size=9, bold=False, color=BLACK, fill=None, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p_style(p, 1, 1)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, text, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_border(cell)
    for c in cell.paragraphs:
        for run in c.runs:
            run.font.name = "Calibri"


def header_row(table, labels, fill="1A365D"):
    row = table.rows[0]
    for i, lab in enumerate(labels):
        cell_text(row.cells[i], lab, size=8, bold=True, color=WHITE, fill=fill, align="center")


def blank_rows(table, start, n, cols):
    for r in range(start, start + n):
        for c in range(cols):
            cell_text(table.rows[r].cells[c], " ", size=9)


def section_bar(doc, title):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_text(t.cell(0, 0), title, size=11, bold=True, color=WHITE, fill="1A365D", align="center")
    doc.add_paragraph()


def field_table(doc, pairs):
    """pairs: list of (label, value) spanning 2 columns (label|value|label|value)."""
    n = (len(pairs) + 1) // 2
    tbl = doc.add_table(rows=n, cols=4)
    tbl.autofit = True
    i = 0
    for r in range(n):
        for c_off in (0, 2):
            if i >= len(pairs):
                cell_text(tbl.rows[r].cells[c_off], " ", size=8, fill="E8EEF6")
                cell_text(tbl.rows[r].cells[c_off + 1], " ", size=9)
                continue
            lab, val = pairs[i]
            cell_text(tbl.rows[r].cells[c_off], lab, size=8, bold=True, fill="E8EEF6")
            cell_text(tbl.rows[r].cells[c_off + 1], val, size=9)
            i += 1
    return tbl


def note(doc, text, size=8, italic=True):
    p = doc.add_paragraph()
    p_style(p, 2, 6)
    add_text(p, text, size=size, bold=False, color=MUTED)
    if italic:
        p.runs[0].italic = True


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.4)
    sec.right_margin = Cm(1.4)
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.2)

    h = doc.add_paragraph()
    p_style(h, 0, 0, WD_ALIGN_PARAGRAPH.CENTER)
    add_text(h, "INSTITUTO NACIONAL DO SEGURO SOCIAL - INSS", size=11, bold=True, color=NAVY)

    h2 = doc.add_paragraph()
    p_style(h2, 0, 2, WD_ALIGN_PARAGRAPH.CENTER)
    add_text(h2, "PERFIL PROFISSIOGRAFICO PREVIDENCIARIO - PPP", size=16, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    p_style(sub, 0, 8, WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        sub,
        "Modelo atual - Anexo XVII da IN PRES/INSS n. 128/2022 (alteracoes da IN n. 133/2022)",
        size=9,
        color=MUTED,
    )

    aviso = doc.add_table(rows=1, cols=1)
    cell = aviso.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p_style(p, 4, 4)
    add_text(p, "USO: passar a limpo PPP em papel (ex.: vinculo de 2017) e assinar de novo. ", size=9, bold=True)
    add_text(
        p,
        "Empresa, trabalhador e exposicao ficam em branco. Preencher com os dados do PPP original e do LTCAT/PGR/PCMSO. "
        "Periodos ate 31/12/2022: PPP em papel emitido pela empresa. Periodos a partir de 01/01/2023: PPP eletronico (PPP-e) via eSocial / Meu INSS. "
        "Este arquivo nao substitui o PPP-e nem o laudo tecnico. Informacao falsa e crime (art. 297 do Codigo Penal).",
        size=9,
    )
    shade_cell(cell, "FFF4E5")
    set_cell_border(cell, "C05621", "12")

    doc.add_paragraph()
    section_bar(doc, "I - SECAO DE DADOS ADMINISTRATIVOS")

    emp = doc.add_paragraph()
    p_style(emp, 0, 4)
    add_text(emp, "EMPREGADOR / ESTABELECIMENTO  (preencher a empresa do vinculo, nao a PerfectGest)", size=9, bold=True, color=NAVY)

    field_table(
        doc,
        [
            ("1. CNPJ / CEI / CAEPF / CNO", LINE[:40]),
            ("3. CNAE", LINE[:28]),
            ("2. Nome empresarial (razao social)", LINE),
            ("Nome fantasia (se houver)", LINE[:40]),
            ("Endereco (logradouro, n., bairro)", LINE),
            ("Municipio / UF / CEP", LINE[:40]),
            ("Telefone", LINE[:28]),
            ("E-mail", LINE[:40]),
        ],
    )
    note(doc, "O item 1 e o CNPJ do empregador. No item 13.2 informe o CNPJ do estabelecimento onde o trabalhador estava lotado (tomadora, se terceirizado).")

    trab = doc.add_paragraph()
    p_style(trab, 6, 4)
    add_text(trab, "TRABALHADOR", size=9, bold=True, color=NAVY)

    field_table(
        doc,
        [
            ("4. Nome do trabalhador", LINE),
            ("5. BR / PDH", "[  ] BR    [  ] PDH"),
            ("6. CPF  (modelo atual; no PPP antigo era NIT)", LINE[:36]),
            ("7. Data de nascimento", "____ / ____ / ________"),
            ("8. Sexo (M/F)", "________"),
            ("9. Matricula eSocial  (no PPP antigo: CTPS)", LINE[:36]),
            ("10. Data de admissao", "____ / ____ / ________"),
            ("11. Regime de revezamento  (ou NA)", LINE[:36]),
            ("12.1 Data da CAT (se houver)", "____ / ____ / ________"),
            ("12.2 Numero da CAT", LINE[:28]),
        ],
    )
    note(
        doc,
        "Item 5: BR = brasileiro nato/naturalizado; PDH = portador de deficiencia habilitado. "
        "Item 9: se o periodo for anterior ao eSocial, anote a CTPS (n. / serie / UF) nas Observacoes e deixe a matricula se nao existir.",
    )

    lot = doc.add_paragraph()
    p_style(lot, 6, 4)
    add_text(lot, "13. LOTACAO E ATRIBUICAO  (uma linha por periodo; nova linha se cargo, setor ou codigo GFIP mudar)", size=9, bold=True, color=NAVY)

    t13 = doc.add_table(rows=5, cols=7)
    header_row(
        t13,
        [
            "13.1 Periodo\n(de / ate)",
            "13.2 CNPJ/CEI/\nCAEPF/CNO",
            "13.3 Setor",
            "13.4 Cargo",
            "13.5 Funcao\n(ou NA)",
            "13.6 CBO",
            "13.7 Codigo GFIP\n/ eSocial",
        ],
    )
    blank_rows(t13, 1, 4, 7)
    note(
        doc,
        "13.7 (unico vinculo): em branco = nunca exposto; 01 = ja esteve exposto, nao neste periodo; 02/03/04 = exposicao 15/20/25 anos. "
        "Mais de um vinculo no periodo: 05 a 08. Transcrever o codigo declarado na GFIP/eSocial da epoca.",
    )

    prof = doc.add_paragraph()
    p_style(prof, 6, 4)
    add_text(prof, "14. PROFISSIOGRAFIA  (descrever as atividades reais, nao so o cargo da CTPS)", size=9, bold=True, color=NAVY)

    t14 = doc.add_table(rows=5, cols=2)
    header_row(t14, ["14.1 Periodo (de / ate)", "14.2 Descricao das atividades"])
    blank_rows(t14, 1, 4, 2)
    t14.rows[0].cells[1].width = Cm(14)
    note(doc, "Preferir os mesmos periodos do item 13. Texto deve bater com o LTCAT.")

    section_bar(doc, "II - SECAO DE REGISTROS AMBIENTAIS")

    exp = doc.add_paragraph()
    p_style(exp, 0, 4)
    add_text(exp, "15. EXPOSICAO A FATORES DE RISCOS", size=9, bold=True, color=NAVY)

    t15 = doc.add_table(rows=6, cols=8)
    header_row(
        t15,
        [
            "15.1 Periodo",
            "15.2 Tipo",
            "15.3 Fator de risco",
            "15.4 Intensidade /\nconcentracao",
            "15.5 Tecnica",
            "15.6 EPC\neficaz S/N",
            "15.7 EPI\neficaz S/N",
            "15.8 CA EPI",
        ],
    )
    blank_rows(t15, 1, 5, 8)
    note(
        doc,
        "15.2 Tipo: 1 Fisico | 2 Quimico | 3 Biologico | 4 Ergonomico/psicossocial (facultativo) | 5 Mecanico/acidente (facultativo). "
        "Ruido: dB(A), metodologia NHO-01 FUNDACENTRO ou NR-15 (indicar a norma). Se nao houver agente: NA.",
    )

    epi = doc.add_paragraph()
    p_style(epi, 6, 4)
    add_text(epi, "15.9 ATENDIMENTO NR-06 e NR-09 (EPI informado)  - responder S / N / NA", size=9, bold=True, color=NAVY)

    t159 = doc.add_table(rows=6, cols=2)
    header_row(t159, ["Item", "Pergunta / resposta"])
    qs = [
        (
            "15.9.1 Medida de protecao",
            "Foi tentada implementacao de medida coletiva, administrativa ou de organizacao do trabalho, optando-se pelo EPI por inviabilidade tecnica, insuficiencia, interinidade, ou em caracter complementar/emergencial?    S [  ]   N [  ]   NA [  ]",
        ),
        (
            "15.9.2 Funcionamento do EPI",
            "Foram observadas as condicoes de funcionamento e do uso ininterrupto do EPI ao longo do tempo, conforme especificacao do fabricante, ajustada as condicoes de campo?    S [  ]   N [  ]   NA [  ]",
        ),
        (
            "15.9.3 Prazo de validade",
            "Foi observado o prazo de validade conforme Certificado de Aprovacao (CA) do MTE/MTP?    S [  ]   N [  ]   NA [  ]",
        ),
        (
            "15.9.4 Periodicidade de troca",
            "Foi observada a periodicidade de troca definida pelos programas ambientais, comprovada por recibo assinado pelo usuario em epoca propria?    S [  ]   N [  ]   NA [  ]",
        ),
        (
            "15.9.5 Higienizacao",
            "Foi observada a higienizacao do EPI?    S [  ]   N [  ]   NA [  ]",
        ),
    ]
    for i, (lab, q) in enumerate(qs, start=1):
        cell_text(t159.rows[i].cells[0], lab, size=8, bold=True, fill="E8EEF6")
        cell_text(t159.rows[i].cells[1], q, size=8)

    bio = doc.add_paragraph()
    p_style(bio, 10, 4)
    add_text(bio, "MONITORACAO BIOLOGICA  (se houver PCMSO/exames; senao NA)", size=9, bold=True, color=NAVY)

    tbio = doc.add_table(rows=4, cols=5)
    header_row(tbio, ["Periodo", "Exame / indicador", "Data", "Resultado", "Observacao"])
    blank_rows(tbio, 1, 3, 5)

    section_bar(doc, "III - RESPONSAVEL PELOS REGISTROS AMBIENTAIS")

    ra = doc.add_paragraph()
    p_style(ra, 0, 4)
    add_text(
        ra,
        "16. Profissional legalmente habilitado (medico do trabalho ou engenheiro de seguranca do trabalho). Sem este bloco o PPP nao prova especialidade.",
        size=9,
        bold=True,
        color=NAVY,
    )

    t16 = doc.add_table(rows=4, cols=4)
    header_row(t16, ["16.1 Periodo", "16.2 CPF", "16.3 Registro conselho (CRM/CREA)", "16.4 Nome do profissional"])
    blank_rows(t16, 1, 3, 4)
    note(doc, "Modelo atual pede CPF (nao NIT). Transcrever do PPP/LTCAT original. Nao inventar registro profissional.")

    section_bar(doc, "IV - RESPONSAVEIS PELAS INFORMACOES")

    decl = doc.add_table(rows=1, cols=1)
    dcell = decl.cell(0, 0)
    dcell.text = ""
    dp = dcell.paragraphs[0]
    p_style(dp, 4, 4)
    add_text(
        dp,
        "Declaramos, para todos os fins de direito, que as informacoes prestadas neste documento sao veridicas e foram transcritas fielmente dos registros administrativos, das demonstracoes ambientais e dos programas medicos de responsabilidade da empresa. "
        "E de nosso conhecimento que a prestacao de informacoes falsas neste documento constitui crime de falsificacao de documento publico, nos termos do artigo 297 do Codigo Penal e, tambem, que tais informacoes sao de caracter privativo do trabalhador, constituindo crime, nos termos da Lei n. 9.029/95, praticas discriminatorias decorrentes de sua exigibilidade por outrem, bem como de sua divulgacao para terceiros, ressalvado quando exigida pelos orgaos publicos competentes.",
        size=8,
    )
    shade_cell(dcell, "F3F4F6")
    set_cell_border(dcell)

    emi = doc.add_paragraph()
    p_style(emi, 10, 4)
    add_text(emi, "17. DATA DE EMISSAO DO PPP:  ____ / ____ / ________", size=10, bold=True)

    rl = doc.add_paragraph()
    p_style(rl, 8, 4)
    add_text(rl, "18. REPRESENTANTE LEGAL DA EMPRESA  (quem assina pelo empregador do vinculo)", size=9, bold=True, color=NAVY)

    t18 = doc.add_table(rows=2, cols=2)
    header_row(t18, ["18.1 CPF do representante legal", "18.2 Nome do representante legal"])
    cell_text(t18.rows[1].cells[0], LINE[:36], size=10)
    cell_text(t18.rows[1].cells[1], LINE, size=10)
    note(doc, "Modelo atual: CPF (nao NIT). Cargo/funcao na empresa: preencher abaixo. Assinar e carimbar.")

    field_table(
        doc,
        [
            ("Cargo / funcao na empresa", LINE[:40]),
            ("Local da assinatura", LINE[:40]),
        ],
    )

    sig = doc.add_table(rows=2, cols=2)
    cell_text(sig.cell(0, 0), "\n\n\n" + "_" * 42 + "\nAssinatura do representante legal", size=9, align="center")
    cell_text(sig.cell(0, 1), "\n\n\n" + "_" * 42 + "\nCarimbo da empresa (CNPJ)", size=9, align="center")
    cell_text(sig.cell(1, 0), "\n\n\n" + "_" * 42 + "\nAssinatura do profissional item 16 (se distinta)", size=8, align="center")
    cell_text(sig.cell(1, 1), "\n\n\n" + "_" * 42 + "\nCarimbo CRM / CREA", size=8, align="center")

    obs = doc.add_paragraph()
    p_style(obs, 12, 4)
    add_text(obs, "OBSERVACOES", size=9, bold=True, color=NAVY)
    tobs = doc.add_table(rows=4, cols=1)
    cell_text(tobs.cell(0, 0), "Alteracao de razao social, CTPS (se periodo anterior ao eSocial), layout, sucessao, etc.:", size=8, fill="E8EEF6")
    for i in range(1, 4):
        cell_text(tobs.rows[i].cells[0], " ", size=11)

    check = doc.add_paragraph()
    p_style(check, 12, 2)
    add_text(check, "CHECKLIST PARA PASSAR A LIMPO O PPP DE 2017", size=10, bold=True, color=NAVY)
    for item in (
        "[  ] Copiar razao social, CNPJ, CNAE e endereco exatamente como no PPP/contrato da empresa do vinculo.",
        "[  ] Copiar nome, CPF (ou NIT do original no campo 6, convertendo para CPF se tiver), admissao e CAT.",
        "[  ] Repetir lotacao, cargo, CBO, codigo GFIP e descricao das atividades do original (e do LTCAT).",
        "[  ] Repetir agentes, intensidade, tecnica, EPC/EPI, CA e respostas 15.9.",
        "[  ] Identificar o profissional do item 16 com CPF e conselho (do laudo da epoca).",
        "[  ] Preencher data de emissao ATUAL, seu CPF, nome, assinar e carimbar como representante legal daquela empresa.",
        "[  ] Nao inventar exposicao, CBO, codigo GFIP nem registro de engenheiro/medico.",
    ):
        p = doc.add_paragraph()
        p_style(p, 1, 1)
        add_text(p, item, size=9)

    foot = doc.add_paragraph()
    p_style(foot, 14, 0, WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        foot,
        "Fonte do modelo: Anexo XVII da IN PRES/INSS n. 128/2022. PDF oficial: https://www.gov.br/inss/pt-br/centrais-de-conteudo/publicacoes/outras/ppp.pdf",
        size=8,
        color=MUTED,
    )
    foot2 = doc.add_paragraph()
    p_style(foot2, 0, 0, WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        foot2,
        "Arquivo de trabalho interno (docs do site PerfectGest). Nao e documento da Perfect Gest Desenvolvimento de Software Ltda.",
        size=8,
        color=MUTED,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(str(OUT))


if __name__ == "__main__":
    build()
